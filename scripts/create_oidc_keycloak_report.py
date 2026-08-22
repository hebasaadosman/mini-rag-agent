"""Create the local Keycloak OIDC sandbox validation report."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "OIDC_KEYCLOAK_SANDBOX_REPORT.docx"

NAVY = "0B1F3A"
BLUE = "1F6FEB"
LIGHT_BLUE = "EAF2FF"
LIGHT_GREEN = "E8F5EE"
GREEN = "1A7F37"
GRAY = "52606D"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(15 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run(text)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in (("Title", 25, NAVY), ("Subtitle", 11, GRAY), ("Heading 1", 15, NAVY), ("Heading 2", 11, BLUE)):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name == "Title" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name not in {"Subtitle"}

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("MINI-RAG-AGENT  |  SECURITY VALIDATION")
    header_run.font.name = "Aptos"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor.from_string(GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Local sandbox validation - no merge, push, or deployment performed")
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(GRAY)

    p = doc.add_paragraph()
    p.style = "Title"
    p.add_run("Keycloak OIDC Sandbox\nValidation Report")
    sub = doc.add_paragraph("Provider-neutral BFF authentication | 20 August 2026")
    sub.style = "Subtitle"

    add_callout(
        doc,
        "Outcome:",
        "All planned local end-to-end authentication checks passed after resolving two integration issues: the Keycloak test-realm profile requirement and missing asymmetric-JWT crypto support in the FastAPI image.",
        LIGHT_GREEN,
    )

    add_heading(doc, "1. Scope and design", 1)
    doc.add_paragraph(
        "Keycloak was used only as a local development/integration identity provider. "
        "The application remains OIDC provider-neutral: issuer, endpoints, scopes, claims, and client settings are supplied through configuration rather than Keycloak-specific application logic."
    )
    flow = doc.add_table(rows=1, cols=5)
    flow.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["Angular", "FastAPI BFF", "Keycloak sandbox", "OIDC", "Redis session"]
    for index, label in enumerate(labels):
        cell_text(flow.cell(0, index), label, bold=True, color="FFFFFF")
        shade(flow.cell(0, index), NAVY if index in {0, 1, 4} else BLUE)
        flow.cell(0, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Browser flow: Angular -> FastAPI BFF -> Keycloak sandbox -> OIDC -> FastAPI -> Redis session.")

    add_bullet(doc, "Authorization Code flow with PKCE (S256).")
    add_bullet(doc, "FastAPI performs the code-to-token exchange and validates the signed ID token through JWKS.")
    add_bullet(doc, "Angular receives neither access tokens nor refresh tokens; it uses an opaque HttpOnly BFF session cookie.")
    add_bullet(doc, "Keycloak realm roles are mapped into the configured roles claim and resolved by the protected FastAPI endpoints.")
    add_bullet(doc, "Redis stores server-side session and OIDC transaction state. The session URL is explicit in production and may be derived only in local/test environments.")

    add_heading(doc, "2. Sandbox configuration", 1)
    config = doc.add_table(rows=1, cols=2)
    config.alignment = WD_TABLE_ALIGNMENT.LEFT
    config.style = "Table Grid"
    for i, label in enumerate(("Component", "Sandbox configuration")):
        cell_text(config.cell(0, i), label, bold=True, color="FFFFFF")
        shade(config.cell(0, i), NAVY)
    rows = [
        ("Identity provider", "Keycloak 26.3.2, local realm: mini-rag"),
        ("OIDC client", "mini-rag-bff; public client; standard authorization-code flow; PKCE S256"),
        ("Test identities", "heba.admin (platform_admin) and analyst.user (analyst); sandbox-only accounts"),
        ("TLS", "Self-signed loopback certificate for https://127.0.0.1:8444; test client explicitly trusts only this local certificate"),
        ("FastAPI callback", "http://127.0.0.1:8000/api/v1/auth/callback in local development"),
        ("Session behavior", "Redis-backed opaque session; 3-second idle timeout and 20-second absolute timeout for deterministic sandbox testing"),
    ]
    for left, right in rows:
        cells = config.add_row().cells
        cell_text(cells[0], left, bold=True)
        cell_text(cells[1], right)

    add_heading(doc, "3. End-to-end validation results", 1)
    doc.add_paragraph("Command executed locally: scripts/run_keycloak_oidc_e2e.py")
    results = doc.add_table(rows=1, cols=3)
    results.alignment = WD_TABLE_ALIGNMENT.LEFT
    results.style = "Table Grid"
    for i, label in enumerate(("Validation", "Result", "Evidence")):
        cell_text(results.cell(0, i), label, bold=True, color="FFFFFF")
        shade(results.cell(0, i), NAVY)
    checks = [
        ("Login redirect and authorization-code exchange", "PASSED", "BFF redirected to Keycloak; successful callback returned to the local application."),
        ("ID-token and JWKS signature validation", "PASSED", "Keycloak RS256-signed ID token validated against the configured JWKS endpoint."),
        ("Role mapping and protected endpoint access", "PASSED", "Authenticated principal contained the platform_admin role and accessed /api/v1/auth/me."),
        ("Redis session creation", "PASSED", "Session and CSRF cookies were issued; protected endpoint resolved the principal from server-side session state."),
        ("CSRF rejection", "PASSED", "Logout without the configured CSRF header returned 403."),
        ("Logout and revoked-session behavior", "PASSED", "Logout with the matching CSRF header returned 204; the same session then returned 401."),
        ("Idle session expiry", "PASSED", "A fresh session expired after the configured local idle timeout and returned 401."),
        ("Focused unit tests", "PASSED", "9 tests passed for BFF authentication and session URL resolution."),
    ]
    for name, status, evidence in checks:
        cells = results.add_row().cells
        cell_text(cells[0], name)
        cell_text(cells[1], status, bold=True, color=GREEN)
        shade(cells[1], LIGHT_GREEN)
        cell_text(cells[2], evidence)

    add_heading(doc, "4. Issues found and resolved", 1)
    issues = [
        ("Keycloak profile action interrupted the test login", "The initial realm users lacked profile fields and Keycloak requested profile completion. Added sandbox-only first name, last name, and verified email fields, then re-imported the sandbox realm."),
        ("FastAPI could not validate Keycloak RS256 signatures", "PyJWT was present but its cryptography backend was missing. Added cryptography==46.0.5 to src/requirements.txt. The current local container was restarted after a temporary install for the validation run; future image builds pick it up from requirements."),
        ("Test assumed OIDC subject equals username", "Corrected the test: OIDC sub is an opaque stable identifier, and Keycloak returns a UUID. The test now verifies a non-empty subject plus the expected role."),
        ("HTTPS health check in Keycloak image", "The container lacks curl; local health check now verifies the management TCP port after start-up. Application-facing HTTPS was verified through the real login flow."),
    ]
    for title, body in issues:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title + ": ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(body)

    add_heading(doc, "5. Production prerequisites and next gate", 1)
    add_callout(
        doc,
        "Not performed:",
        "No merge, push, or deployment was performed as part of this sandbox validation.",
    )
    prerequisites = [
        "Integrate a real OIDC tenant (for example Entra ID) through configuration: issuer discovery, client ID, redirect URI, scopes, JWKS, and claim mapping. No architecture redesign is expected.",
        "Use HTTPS end to end with a trusted certificate; set secure cookies and production callback/front-end URLs.",
        "Provide production Redis through a managed/secured endpoint, with an explicit session URL and credentials from secret management.",
        "Store client secrets and signing/configuration values in the deployment secret store, not source files or frontend bundles.",
        "Confirm production role/entitlement claim mapping and authorization policy with the target organization.",
        "Run the same end-to-end suite against the real sandbox tenant before deciding whether to merge or deploy.",
    ]
    for item in prerequisites:
        add_bullet(doc, item)

    add_heading(doc, "6. Files added or updated for the local sandbox", 1)
    for item in [
        "docker/docker-compose.oidc-sandbox.yml",
        "docker/keycloak/mini-rag-realm.json",
        "docker/env/.env.keycloak-sandbox.example and docker/env/.env.oidc-sandbox.example",
        "scripts/run_keycloak_oidc_e2e.py",
        "src/authentication/session_url.py and tests/test_auth_session_url.py",
        "src/requirements.txt (cryptography dependency)",
        "docs/OIDC_KEYCLOAK_SANDBOX.md",
    ]:
        add_bullet(doc, item)

    doc.core_properties.title = "Keycloak OIDC Sandbox Validation Report"
    doc.core_properties.subject = "Local Keycloak end-to-end validation of provider-neutral BFF authentication"
    doc.core_properties.author = "Mini RAG Agent"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
